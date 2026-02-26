from __future__ import annotations

from typing import Dict, List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt

from auto_paper.models import JournalSpec


def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def export_manuscript_docx(
    output_path: str,
    journal: JournalSpec,
    drafts: Dict[str, str],
    references: List[str],
    title: Optional[str] = None,
):
    doc = Document()

    # Title page (minimal)
    if title:
        p = doc.add_paragraph(title)
        p.runs[0].font.size = Pt(16)
        p.runs[0].bold = True
        doc.add_paragraph("")  # spacer

    # Abstract
    if drafts.get("Abstract"):
        _add_heading(doc, "Abstract", level=1)
        doc.add_paragraph(drafts["Abstract"].strip())

    # Keywords (optional)
    # If user wants, can be added from facts; omitted in MVP.

    # Main text sections
    for sec in ["Introduction", "Methods", "Results", "Discussion", "Conclusion"]:
        if drafts.get(sec):
            _add_heading(doc, sec, level=1)
            doc.add_paragraph(drafts[sec].strip())

    # References
    _add_heading(doc, "References", level=1)
    for r in references:
        doc.add_paragraph(r)

    doc.save(output_path)


def export_cover_letter_docx(
    output_path: str,
    journal: JournalSpec,
    cover_letter_text: str,
):
    doc = Document()
    # Minimal layout; users can customize with letterhead later.
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d"))
    doc.add_paragraph("")
    doc.add_paragraph(f"To the Editor-in-Chief, {journal.journal_name}")
    doc.add_paragraph("")
    doc.add_paragraph(cover_letter_text.strip())
    doc.save(output_path)
