from __future__ import annotations

import os
import re
import base64
import mimetypes
from typing import List, Optional, Tuple

import pandas as pd
import pdfplumber
from docx import Document

from auto_paper.models import TableData, FigureData, StudyFactSheet
from auto_paper.utils import normalize_whitespace, clamp_str


def _guess_mime(path: str) -> str:
    mt, _ = mimetypes.guess_type(path)
    return mt or "image/png"


def image_path_to_data_url(path: str) -> str:
    """Encode a local image file into a base64 data URL for OpenAI vision inputs."""
    mime = _guess_mime(path)
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return f"data:{mime};base64,{b64}"


def table_to_markdown(table: TableData, max_rows: int = 35, max_cols: int = 14, cell_char_limit: int = 48) -> str:
    """Render a compact markdown table preview (token-friendly)."""
    def _clamp(s: str) -> str:
        s = normalize_whitespace(str(s))
        s = s.replace("|", "\\|")
        if len(s) > cell_char_limit:
            return s[: cell_char_limit - 1] + "…"
        return s

    header = table.header[:max_cols]
    rows = table.rows
    if len(rows) > max_rows:
        head_n = max_rows - 5
        rows = rows[: max(0, head_n)] + rows[-5:]

    header_cells = ["" if h is None else _clamp(h) for h in header]
    out = [
        "| " + " | ".join(header_cells) + " |",
        "| " + " | ".join(["---"] * len(header_cells)) + " |",
    ]
    for r in rows:
        cells = [_clamp(c) for c in r[:max_cols]]
        # pad to header length
        if len(cells) < len(header_cells):
            cells += [""] * (len(header_cells) - len(cells))
        out.append("| " + " | ".join(cells) + " |")

    note = ""
    if len(table.rows) > max_rows:
        note = f"\n\n(Note: table preview truncated; showing first {max_rows-5} rows and last 5 rows.)"
    if len(table.header) > max_cols:
        note += f"\n(Note: columns truncated to first {max_cols}.)"
    return "\n".join(out) + note


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return normalize_whitespace("\n\n".join(paras))


def extract_text_from_pdf(path: str) -> str:
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                texts.append(t)
    return normalize_whitespace("\n\n".join(texts))


def parse_tables_from_docx(path: str, prefix: str = "T") -> List[TableData]:
    doc = Document(path)
    out: List[TableData] = []
    t_idx = 1
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [normalize_whitespace(cell.text) for cell in row.cells]
            rows.append(cells)

        if not rows:
            continue
        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        out.append(TableData(id=f"{prefix}{t_idx}", title=None, header=header, rows=body))
        t_idx += 1
    return out


def parse_tables_from_excel(path: str, sheet_name: Optional[str] = None, prefix: str = "T") -> List[TableData]:
    # read first sheet by default
    xls = pd.ExcelFile(path)
    sheets = [sheet_name] if sheet_name else xls.sheet_names[:1]
    out: List[TableData] = []
    t_idx = 1
    for sh in sheets:
        df = pd.read_excel(path, sheet_name=sh, dtype=str)
        df = df.fillna("")
        if df.shape[0] == 0 or df.shape[1] == 0:
            continue
        header = [str(c).strip() for c in df.columns.tolist()]
        rows = df.astype(str).values.tolist()
        out.append(TableData(id=f"{prefix}{t_idx}", title=sh, header=header, rows=rows))
        t_idx += 1
    return out


def infer_figure_type(filename: str) -> str:
    fn = filename.lower()
    if "km" in fn or "kaplan" in fn:
        return "KaplanMeier"
    if "forest" in fn:
        return "ForestPlot"
    if "bar" in fn:
        return "BarChart"
    return "Other"


def build_fact_sheet_with_llm(
    llm_client,
    plan_text: str,
    tables: List[TableData],
    figures: List[FigureData],
    notes_for_writer: Optional[str] = None,
    enable_vision_for_figures: bool = True,
) -> StudyFactSheet:
    """
    Use LLM to extract a structured fact sheet.
    IMPORTANT: must not invent numeric results; if unknown, use null.
    """
    system = (
        "You are a clinical research methodologist and manuscript editor. "
        "Extract a structured StudyFactSheet JSON. "
        "CRITICAL: Do NOT invent numbers, p-values, confidence intervals, or sample sizes. "
        "Only use numbers explicitly present in the provided plan/tables. "
        "If uncertain, set fields to null. Use concise strings."
    )

    # NOTE: We intentionally do NOT ask the model to echo the full tables/figures.
    # We keep raw tables/figures from the user's uploads (ground truth) and only let the model
    # fill the higher-level study metadata.

    table_signals = [
        {
            "id": t.id,
            "title": t.title,
            "n_rows": len(t.rows),
            "n_cols": len(t.header),
            "header": t.header[: min(12, len(t.header))],
        }
        for t in tables
    ]
    figure_signals = [
        {"id": f.id, "filename": f.filename, "figure_type": f.figure_type, "caption_raw": f.caption_raw}
        for f in figures
    ]

    user = f"""
Plan/Protocol text (truncated):
\"\"\"{clamp_str(plan_text, 14000)}\"\"\"

Available Tables (signals only):
{table_signals}

Available Figures (signals only):
{figure_signals}

Additional author notes (optional):
{notes_for_writer or ""}

Return JSON matching ONLY the following structure (do not include tables/figures):
{{
  "title": null | "string",
  "keywords": ["..."],
  "study_design": null | "string",
  "period": {{"start": "YYYY-MM-DD"|null, "end":"YYYY-MM-DD"|null}},
  "irb": {{"approved": true/false, "number": "string"|null}},
  "population": {{
    "groups": [{{"name":"Group A","n":120,"definition":null}}, ...],
    "inclusion": ["..."],
    "exclusion": ["..."]
  }},
  "endpoints": {{
    "primary": "string"|null,
    "secondary": ["..."]
  }},
  "statistics": {{
    "software": "string"|null,
    "alpha": 0.05,
    "methods": ["..."],
    "notes": "string"|null
  }},
  "protocol_key_points": ["..."],
  "notes_for_writer": "string"|null
}}

Rules:
- Do NOT invent any numeric values. If uncertain, set null.
- Keep keywords short (3–8).
- protocol_key_points: 8–18 bullets summarizing essential protocol/method details (setting, population, intervention/comparator, endpoints, follow-up, stats). No numbers unless explicitly stated.
"""

    meta = llm_client.chat_json(system=system, user=user, temperature=0.0, max_tokens=1800)

    # Merge: keep the original uploaded tables/figures as ground truth.
    facts = StudyFactSheet.model_validate({
        **meta,
        "tables": [t.model_dump() for t in tables],
        "figures": [f.model_dump() for f in figures],
        "plan_text": plan_text,
        "notes_for_writer": meta.get("notes_for_writer") or notes_for_writer,
    })

    # Enrich tables with category + key_findings (per-table JSON extraction).
    if tables:
        facts.tables = enrich_tables_with_llm(llm_client, facts, max_tables=12)

    # Enrich figures with key_points (optionally via vision).
    if figures:
        facts.figures = enrich_figures_with_llm(llm_client, facts, enable_vision=enable_vision_for_figures)

    return facts


def enrich_tables_with_llm(
    llm_client,
    facts: StudyFactSheet,
    max_tables: int = 12,
    max_rows_md: int = 40,
    max_cols_md: int = 14,
) -> List[TableData]:
    """Fill TableData.category and TableData.key_findings while preserving original table cells."""
    system = (
        "You are a clinical biostatistician and manuscript results editor. "
        "Given ONE results table, extract (1) a table category, (2) a clean title, "
        "and (3) 3–8 key findings that are explicitly supported by the table cells. "
        "CRITICAL: do NOT invent numbers, p-values, HR, CI, percentages, or Ns. "
        "Use exact strings from the table when possible. If unclear, return empty key_findings." 
        "Output ONLY valid JSON."
    )

    out: List[TableData] = []
    for t in facts.tables[:max_tables]:
        md = table_to_markdown(t, max_rows=max_rows_md, max_cols=max_cols_md)
        user = f"""
Table ID: {t.id}
Current title: {t.title or ""}

Study context (for interpreting group labels only; do not invent results):
- Study design: {facts.study_design or ""}
- Primary endpoint: {facts.endpoints.primary or ""}
- Groups: {[{"name": g.name, "n": g.n} for g in facts.population.groups]}

Table (markdown preview):
{md}

Return JSON:
{{
  "id": "{t.id}",
  "title": "string"|null,
  "category": "Baseline|Perioperative|Pathologic|Oncologic|Subgroup|Other"|null,
  "key_findings": [
    {{
      "statement": "string",
      "values": {{"GroupA": "...", "GroupB": "..."}},
      "p_value": "..."|null,
      "ci_95": "..."|null
    }}
  ]
}}

Rules:
- key_findings MUST be supported by the table as shown.
- If no p-value column exists, set p_value to null.
- Keep each statement short and clinically meaningful.
"""
        data = llm_client.chat_json(system=system, user=user, temperature=0.0, max_tokens=1200)
        # Merge back into original
        try:
            t2 = TableData.model_validate({
                **t.model_dump(),
                "title": data.get("title") or t.title,
                "category": data.get("category"),
                "key_findings": data.get("key_findings") or [],
            })
        except Exception:
            t2 = t
        out.append(t2)

    # Keep any remaining tables untouched
    if len(facts.tables) > max_tables:
        out.extend(facts.tables[max_tables:])
    return out


def enrich_figures_with_llm(
    llm_client,
    facts: StudyFactSheet,
    enable_vision: bool = True,
    max_figures: int = 12,
) -> List[FigureData]:
    """Fill FigureData.key_points (and optionally caption) from caption and/or image."""
    system = (
        "You are a clinical research editor. Extract 3–8 key points from ONE figure. "
        "CRITICAL: do NOT guess numeric values. Only report numbers that are explicitly printed on the figure. "
        "If the figure contains Kaplan–Meier/forest plot, summarize direction and statistically significant findings only if clearly labeled. "
        "Output ONLY valid JSON."
    )

    out: List[FigureData] = []
    for f in facts.figures[:max_figures]:
        caption = f.caption_raw or ""
        # If we can, send the image via multimodal message blocks (Chat Completions format).
        if enable_vision and f.path and os.path.exists(f.path):
            data_url = image_path_to_data_url(f.path)
            user_blocks = [
                {"type": "text", "text": (
                    f"Figure ID: {f.id}\n"
                    f"Figure type: {f.figure_type}\n"
                    f"Current caption (may be empty): {caption}\n\n"
                    "Task: read the figure and extract key points (no hallucination).\n"
                    "Return JSON: {id, caption_raw, key_points}.\n"
                    "- key_points: 3–8 bullets\n"
                    "- caption_raw: if you can infer a concise caption from the figure labels, write it; otherwise keep the provided caption.\n"
                    "- Do NOT include any extra keys."
                )},
                {"type": "image_url", "image_url": {"url": data_url}},
            ]
            data = llm_client.chat_json(system=system, user=user_blocks, temperature=0.0, max_tokens=900)
        else:
            # Caption-only fallback
            user = f"""
Figure ID: {f.id}
Figure type: {f.figure_type}
Filename: {f.filename}
Caption (may be empty): {caption}

Return JSON:
{{
  "id": "{f.id}",
  "caption_raw": "string"|null,
  "key_points": ["..."]
}}
"""
            data = llm_client.chat_json(system=system, user=user, temperature=0.0, max_tokens=600)

        try:
            f2 = FigureData.model_validate({
                **f.model_dump(),
                "caption_raw": data.get("caption_raw") or f.caption_raw,
                "key_points": data.get("key_points") or [],
            })
        except Exception:
            f2 = f
        out.append(f2)

    if len(facts.figures) > max_figures:
        out.extend(facts.figures[max_figures:])
    return out
